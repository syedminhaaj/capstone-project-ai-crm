from docx import Document
from docx.shared import Pt
from docx.enum.text import WD_ALIGN_PARAGRAPH
import tempfile

def generate_student_template() -> str:
    """Generate a Word document template for student information"""
    
    doc = Document()
    
    # Add title
    title = doc.add_heading('Student Information Template', 0)
    title.alignment = WD_ALIGN_PARAGRAPH.CENTER
    
    # Add instructions
    intro = doc.add_paragraph()
    intro.add_run('Instructions: ').bold = True
    intro.add_run('Fill out this form completely and return it to the driving school office. All fields marked with (*) are required.')
    
    doc.add_paragraph()
    
    # Personal Information Section
    doc.add_heading('Personal Information', level=1)
    
    fields = [
        ('Full Name', '*'),
        ('Date of Birth', '* (MM/DD/YYYY)'),
        ('Email Address', '*'),
        ('Phone Number', '* (XXX) XXX-XXXX'),
        ('Current Address', '*'),
        ('City', '*'),
        ('State', '*'),
        ('ZIP Code', '*'),
    ]
    
    for field, note in fields:
        p = doc.add_paragraph()
        p.add_run(f'{field} {note}: ').bold = True
        p.add_run('_' * 60)
    
    doc.add_paragraph()
    
    # License Information Section
    doc.add_heading('License Information', level=1)
    
    license_fields = [
        ('Driver\'s License Number', '*'),
        ('License State', '*'),
        ('License Expiration Date', '(MM/DD/YYYY)'),
    ]
    
    for field, note in license_fields:
        p = doc.add_paragraph()
        p.add_run(f'{field} {note}: ').bold = True
        p.add_run('_' * 60)
    
    doc.add_paragraph()
    
    # Emergency Contact Section
    doc.add_heading('Emergency Contact Information', level=1)
    
    emergency_fields = [
        ('Emergency Contact Name', '*'),
        ('Relationship', '*'),
        ('Emergency Phone Number', '* (XXX) XXX-XXXX'),
        ('Emergency Email', ''),
    ]
    
    for field, note in emergency_fields:
        p = doc.add_paragraph()
        p.add_run(f'{field} {note}: ').bold = True
        p.add_run('_' * 60)
    
    doc.add_paragraph()
    
    # Additional Information Section
    doc.add_heading('Additional Information', level=1)
    
    additional_fields = [
        ('Preferred Lesson Times', '(e.g., Weekdays, Weekends, Mornings, Afternoons)'),
        ('Any Medical Conditions We Should Know', '(Optional)'),
        ('How Did You Hear About Us?', ''),
    ]
    
    for field, note in additional_fields:
        p = doc.add_paragraph()
        p.add_run(f'{field} {note}: ').bold = True
        p.add_run('_' * 60)
    
    doc.add_paragraph()
    
    # Signature Section
    doc.add_heading('Signature', level=1)
    
    doc.add_paragraph()
    p = doc.add_paragraph()
    p.add_run('Student Signature: ').bold = True
    p.add_run('_' * 40)
    p.add_run('  Date: ').bold = True
    p.add_run('_' * 20)
    
    doc.add_paragraph()
    p = doc.add_paragraph()
    p.add_run('Parent/Guardian Signature (if under 18): ').bold = True
    p.add_run('_' * 30)
    
    # Footer
    doc.add_paragraph()
    footer = doc.add_paragraph('Thank you for choosing DriveSchool Pro!')
    footer.alignment = WD_ALIGN_PARAGRAPH.CENTER
    footer_run = footer.runs[0]
    footer_run.italic = True
    
    # Save to temporary file
    temp_file = tempfile.NamedTemporaryFile(delete=False, suffix='.docx')
    doc.save(temp_file.name)
    temp_file.close()
    
    return temp_file.name